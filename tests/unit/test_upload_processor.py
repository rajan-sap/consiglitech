"""
Unit tests — Upload Processor

Tests in-memory PDF and DOCX parsing, chunking, and metadata attachment.
No external APIs or disk writes required.
"""

import io
import pytest

pytestmark = pytest.mark.unit


# ── Fixtures ────────────────────────────────────────────────────────────────

@pytest.fixture
def sample_pdf_bytes():
    """Create a minimal valid PDF with two pages of text."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(200, 10, text="Page one content about revenue growth.")
    pdf.add_page()
    pdf.cell(200, 10, text="Page two discusses risk factors and outlook.")
    return pdf.output()


@pytest.fixture
def sample_docx_bytes():
    """Create a minimal DOCX in memory."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("This document covers annual financial results.")
    doc.add_paragraph("Total revenue was $50 billion for the fiscal year.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def empty_docx_bytes():
    """DOCX with no text content."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF Parsing ─────────────────────────────────────────────────────────────

class TestParseUploadedPdf:

    def test_extracts_text_from_pages(self, sample_pdf_bytes):
        from ingestion.upload_processor import parse_uploaded_pdf

        docs = parse_uploaded_pdf(sample_pdf_bytes)
        assert len(docs) == 2
        assert "revenue" in docs[0].page_content.lower()
        assert "risk" in docs[1].page_content.lower()

    def test_page_numbers_are_sequential(self, sample_pdf_bytes):
        from ingestion.upload_processor import parse_uploaded_pdf

        docs = parse_uploaded_pdf(sample_pdf_bytes)
        pages = [d.metadata["page_number"] for d in docs]
        assert pages == [1, 2]

    def test_empty_pdf_returns_empty_list(self):
        """A PDF with no extractable text returns an empty list."""
        from fpdf import FPDF
        from ingestion.upload_processor import parse_uploaded_pdf

        pdf = FPDF()
        pdf.add_page()  # blank page, no text
        docs = parse_uploaded_pdf(pdf.output())
        assert docs == []


# ── DOCX Parsing ────────────────────────────────────────────────────────────

class TestParseUploadedDocx:

    def test_extracts_paragraphs(self, sample_docx_bytes):
        from ingestion.upload_processor import parse_uploaded_docx

        docs = parse_uploaded_docx(sample_docx_bytes)
        assert len(docs) == 1
        assert "revenue" in docs[0].page_content.lower()

    def test_empty_docx_returns_empty_list(self, empty_docx_bytes):
        from ingestion.upload_processor import parse_uploaded_docx

        docs = parse_uploaded_docx(empty_docx_bytes)
        assert docs == []


# ── Chunking ────────────────────────────────────────────────────────────────

class TestChunkUploadedDocuments:

    def test_chunks_have_correct_metadata(self, sample_docx_bytes):
        from ingestion.upload_processor import parse_uploaded_docx, chunk_uploaded_documents

        docs = parse_uploaded_docx(sample_docx_bytes)
        chunks = chunk_uploaded_documents(docs, "report.docx")

        assert len(chunks) >= 1
        for chunk in chunks:
            assert chunk.metadata["file_name"] == "report.docx"
            assert chunk.metadata["document_type"] == "User Upload"
            assert "page_number" in chunk.metadata

    def test_long_text_produces_multiple_chunks(self):
        from langchain_core.documents import Document
        from ingestion.upload_processor import chunk_uploaded_documents

        long_text = "Financial data. " * 500  # ~8000 chars
        docs = [Document(page_content=long_text, metadata={"page_number": 1})]
        chunks = chunk_uploaded_documents(docs, "big_file.pdf")

        assert len(chunks) > 1

    def test_empty_input_returns_empty(self):
        from ingestion.upload_processor import chunk_uploaded_documents

        chunks = chunk_uploaded_documents([], "empty.pdf")
        assert chunks == []


# ── End-to-End Dispatch ─────────────────────────────────────────────────────

class TestProcessUploadedFile:

    def test_pdf_routing(self, sample_pdf_bytes):
        from ingestion.upload_processor import process_uploaded_file

        chunks = process_uploaded_file("annual_report.pdf", sample_pdf_bytes)
        assert len(chunks) >= 1
        assert all(c.metadata["file_name"] == "annual_report.pdf" for c in chunks)

    def test_docx_routing(self, sample_docx_bytes):
        from ingestion.upload_processor import process_uploaded_file

        chunks = process_uploaded_file("summary.docx", sample_docx_bytes)
        assert len(chunks) >= 1

    def test_unsupported_extension_returns_empty(self):
        from ingestion.upload_processor import process_uploaded_file

        chunks = process_uploaded_file("image.png", b"fake png bytes")
        assert chunks == []

    def test_no_extension_returns_empty(self):
        from ingestion.upload_processor import process_uploaded_file

        chunks = process_uploaded_file("README", b"some text")
        assert chunks == []
