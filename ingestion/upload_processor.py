"""
Upload Processor — Parse and chunk user-uploaded files in memory.

Uses pdfplumber (PDF) and python-docx (DOCX) so no external API keys
are needed and everything works from raw bytes (no temp files).
"""

import io
from typing import List

import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from ingestion.constants import ANNUAL_REPORT_SPLITTER


# ── Parsers ─────────────────────────────────────────────────────────────────

def parse_uploaded_pdf(file_bytes: bytes) -> List[Document]:
    """Extract text page-by-page from a PDF using pdfplumber."""
    docs = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"page_number": i},
                ))
    return docs


def parse_uploaded_docx(file_bytes: bytes) -> List[Document]:
    """Extract text from a DOCX file."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not full_text.strip():
        return []
    return [Document(page_content=full_text, metadata={"page_number": 1})]


# ── Chunking ────────────────────────────────────────────────────────────────

def chunk_uploaded_documents(docs: List[Document], file_name: str) -> List[Document]:
    """Split parsed documents into chunks with metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=ANNUAL_REPORT_SPLITTER["chunk_size"],
        chunk_overlap=ANNUAL_REPORT_SPLITTER["chunk_overlap"],
        separators=ANNUAL_REPORT_SPLITTER["separators"],
    )

    chunks = []
    for doc in docs:
        splits = splitter.split_text(doc.page_content)
        for split in splits:
            chunks.append(Document(
                page_content=split,
                metadata={
                    "file_name": file_name,
                    "page_number": doc.metadata.get("page_number", 1),
                    "document_type": "User Upload",
                },
            ))
    return chunks


# ── Public entry point ──────────────────────────────────────────────────────

def process_uploaded_file(file_name: str, file_bytes: bytes) -> List[Document]:
    """Parse and chunk an uploaded file. Returns empty list on unsupported type."""
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    if ext == "pdf":
        docs = parse_uploaded_pdf(file_bytes)
    elif ext == "docx":
        docs = parse_uploaded_docx(file_bytes)
    else:
        return []

    return chunk_uploaded_documents(docs, file_name)
