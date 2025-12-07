"""
Document Loading and Chunk Creation Utilities.

Handles:
- Loading PDF, DOCX, TXT files
- Creating metadata-rich Document chunks
- File path helpers
"""

import os
from typing import List, Optional

from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain_core.documents import Document
from llama_parse import LlamaParse

from .constants import DOC_TYPE_ANNUAL_REPORT, DOC_TYPE_NEWS_ARTICLE

load_dotenv()


# =============================================================================
# DOCUMENT LOADERS
# =============================================================================


def load_pdf(file_path: str) -> List[Document]:
    """
    Extract text from PDF using LlamaParse.
    
    Returns one Document per page with page number in metadata.
    Filters out empty pages and placeholder content.
    """
    parser = LlamaParse(
        result_type="markdown",
        parsing_instruction=(
            "Only use markdown heading syntax (# or ##) for text that is clearly "
            "a section heading or title in the document structure. "
            "Regular paragraph text should remain as plain text without markdown formatting, "
            "even if it appears bold or in a larger font."
        ),
    )
    parsed_docs = parser.load_data(file_path)

    if not parsed_docs:
        return []

    # Placeholder/empty content patterns to filter out
    skip_patterns = ["NO_CONTENT_HERE", "NO CONTENT", "[BLANK PAGE]", ""]

    # Create one Document per page with correct page number
    documents = []
    for i, doc in enumerate(parsed_docs, start=1):
        text = doc.text.strip()
        
        # Skip empty or placeholder content
        if not text or text.upper() in [p.upper() for p in skip_patterns]:
            continue
            
        documents.append(Document(
            page_content=text,
            metadata={"page_number": i}
        ))
    
    return documents


def load_docx(file_path: str) -> List[Document]:
    """Extract text from DOCX file."""
    doc = DocxDocument(file_path)
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if not text.strip():
        return []

    return [Document(page_content=text, metadata={"page_number": 1})]


def load_txt(file_path: str) -> List[Document]:
    """Extract text from TXT file."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    if not text.strip():
        return []

    return [Document(page_content=text, metadata={"page_number": 1})]


def load_document(file_path: str) -> List[Document]:
    """
    Load any supported document type.
    
    Auto-selects loader based on file extension.
    Supported: .pdf, .docx, .txt
    """
    ext = os.path.splitext(file_path)[1].lower()

    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".txt": load_txt,
    }

    if ext not in loaders:
        raise ValueError(f"Unsupported file type: {ext}")

    return loaders[ext](file_path)


# =============================================================================
# CHUNK CREATORS
# =============================================================================


def create_report_chunk(
    text: str,
    file_name: str,
    company: str,
    year: str,
    page_number: int = 1,
) -> Document:
    """Create a chunk with Annual Report metadata."""
    return Document(
        page_content=text,
        metadata={
            "file_name": file_name,
            "document_type": DOC_TYPE_ANNUAL_REPORT,
            "company": company,
            "year": year,
            "page_number": page_number,
        },
    )


def create_news_chunk(
    text: str,
    file_name: str,
    page_number: int = 1,
) -> Document:
    """Create a chunk with News Article metadata."""
    return Document(
        page_content=text,
        metadata={
            "file_name": file_name,
            "document_type": DOC_TYPE_NEWS_ARTICLE,
            "page_number": page_number,
        },
    )


# =============================================================================
# HELPERS
# =============================================================================


def extract_year_from_filename(file_name: str) -> str:
    """
    Extract year from filename like 'report_2023.pdf'.
    
    Returns 'Unknown' if no year found.
    """
    try:
        year = file_name.split("_")[-1].split(".")[0]
        if year.isdigit() and len(year) == 4:
            return year
    except (IndexError, ValueError):
        pass
    return "Unknown"


def get_company_from_path(file_path: str, company_folders: List[str]) -> Optional[str]:
    """
    Extract company name from file path.
    
    Returns None if file is not in a company folder.
    """
    path_parts = file_path.replace("\\", "/").split("/")
    for part in path_parts:
        if part in company_folders:
            return part
    return None
